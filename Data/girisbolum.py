import os
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.text.paragraph import Paragraph
def Iceriyomu(dosyayol):
    document = Document('{}'.format(dosyayol))
    headings = []
    texts = []
    para = []
    giris = ""
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            if headings:
                texts.append(para)
            headings.append(paragraph.text)
            para = []
        elif paragraph.style.name == "Normal" and not paragraph.text.find(' ',0,1) != -1 and paragraph.text !='':
            para.append(paragraph.text)
    if para or len(headings)>len(texts):
        texts.append(texts.append(para))

    for h, t in zip(headings, texts):
        if h== "GİRİŞ" or h== "Giriş":
            giris = t[-1]
            



    if (giris.find('kapsam') != -1 or giris.find('organizasyon') != -1): 
        sonuc="Giris bölümünün son bölümünde tezin organizasyonu ve kapsamına yer verilmis "
        RaporaEkle(sonuc)
    else:
        sonuc="Giris bölümünün son bölümünde tezin organizasyonu ve kapsamına yer verilmemis"
        RaporaEkle(sonuc)

def RaporaEkle(sonuc):
    
    f = open('WordRapor.docx', 'rb')

    document = Document(f)

    
    document.add_paragraph(
        sonuc, style='List Number'
        )

    document.add_heading('16542509-Harun Kurt', level=1)
    document.add_heading('175541018-Yusuf Çelik', level=1)
   
    document.save('WordRapor.docx')

    f.close()

    print("İkinci aşama tamamlandi...")
    print("Word Raporu Olusturuldu...")
