import docx
import os
from docx import Document
from docx.text.paragraph import Paragraph

def Paragrafmidegilmi(dosyayol):
    doc = docx.Document('{}'.format(dosyayol))
    content = doc.paragraphs
    paragraf="1"
    paragrafdegil="0"
    liste=[]
    
    for p in content:

    
        if(p.paragraph_format.first_line_indent == None):
            liste.append(paragrafdegil)
            
        else:
            liste.append(paragraf)
            
    for i in range(len(liste)-1):
        if(liste[i]==liste[i+1] and liste[i]=="1"):
            sonuc="İki satırdan az paragraf var"
            RaporaEkle(sonuc)
            break
        elif((len(liste)-1)=="1"):
            sonuc="İki satırdan az paragraf var"
            RaporaEkle(sonuc)
            break

def RaporaEkle(sonuc):
    
    f = open('WordRapor.docx', 'rb')

    document = Document(f)

    
    document.add_paragraph(
        sonuc, style='List Number'
        )


    document.save('WordRapor.docx')

    f.close()
    print("Asama iki tamamlandi...")
    print("islem Word Raporuna eklendi...")
